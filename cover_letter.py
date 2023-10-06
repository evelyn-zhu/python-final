from docx import Document
from similarity import extract_skills, compute_similarity_and_skills, HARD_SKILLS, preprocess_text
from webscraper import get_job_data


def generate_cover_letter(row_number, resume_text):
    # Load the data
    df = get_job_data()

    if row_number not in df.index:
        print(f"Row number {row_number} not found!")
        return

    # Extract job details
    position_name = df.loc[row_number, "Position Name"]
    company_info = df.loc[row_number, "Company"]  # Get the company name

    # Extract overlapping skills
    processed_resume = preprocess_text(resume_text)
    similarity_results = compute_similarity_and_skills(df, processed_resume)
    if row_number >= len(similarity_results):
        print(f"No data available for row number {row_number}.")
        return
    overlapping_skills = similarity_results[row_number]["Overlapping Skills"]

    # Create a new Document for the cover letter
    doc = Document()

    # Add content based on a simple template structure
    doc.add_paragraph(f"Dear Hiring Manager at {company_info},")
    doc.add_paragraph(f"\nI am writing to express my interest in the {position_name} position listed on LinkedIn. "
                      f"My expertise in {', '.join(overlapping_skills)} makes me a suitable candidate for this role.")

    # Here, you can add more content to your cover letter. For instance:
    doc.add_paragraph(f"\nGiven my experience in these areas, I am confident in my ability "
                      f"to contribute effectively to your team at {company_info}.")

    doc.add_paragraph(
        f"\nThank you for considering my application. I look forward to the opportunity for an interview.")
    doc.add_paragraph(f"\nSincerely,\n[Your Name]")

    # Save the tailored cover letter
    new_filename = f"Cover_Letter_for_{position_name}.docx"
    new_filename = sanitize_filename(new_filename)
    doc.save(new_filename)

    print(f"Cover letter tailored for {position_name} saved as {new_filename}!")

def display_jobs_with_numbers(job_data_df):
    print("\nList of Jobs:")
    for index, row in job_data_df.iterrows():
        if row["Job Description"] != "INVALID LINK":
            print(f"Row {index}: {row['Position Name']} at {row['Company']}")

def sanitize_filename(filename):
    """Sanitize the filename to remove characters that might cause issues."""
    invalid_chars = ["<", ">", ":", "\"", "/", "\\", "|", "?", "*"]
    for char in invalid_chars:
        filename = filename.replace(char, "_")
    return filename
